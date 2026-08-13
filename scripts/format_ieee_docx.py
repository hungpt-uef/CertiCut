"""Build and format the IEEE-style CertiCut Word document."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Cm


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "paper" / "CertiCut_IEEE_Paper.docx"


def _set_font(run, name="Times New Roman", size=9.5, bold=None, italic=None):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _set_columns(section, count, space_twips=360):
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))


def _keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    kwn = pPr.find(qn("w:keepNext"))
    if kwn is None:
        pPr.append(OxmlElement("w:keepNext"))


def _center_table(table):
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")


def _center_inline_shapes(doc):
    """Center every paragraph that contains an inline image."""
    for p in doc.paragraphs:
        if p._p.findall(qn("w:r") + "/" + qn("w:drawing")):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def format_document():
    doc = Document(DOCX)

    # --- Global font ---
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Heading styles ---
    for style_name, size, bold in [
        ("Title", 18, True), ("Subtitle", 11, True),
        ("Heading 1", 10, True), ("Heading 2", 9.5, True),
    ]:
        if style_name in doc.styles:
            s = doc.styles[style_name]
            s.font.name = "Times New Roman"
            s.font.size = Pt(size)
            s.font.bold = bold

    # --- Page layout ---
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    _set_columns(section, 1)

    # --- Title ---
    title_p = doc.paragraphs[0]
    title_p.text = ("CertiCut: Capacitated K-Way Partitioning with Anytime Bounds "
                    "for Sampling-Aware Circuit Cutting")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(8)
    for r in title_p.runs:
        _set_font(r, size=18, bold=True)

    # --- Authors ---
    if len(doc.paragraphs) > 1:
        auth = doc.paragraphs[1]
        auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in auth.runs:
            _set_font(r, size=10)

    # --- Abstract formatting ---
    abs_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Abstract")
    intro_idx = next(i for i, p in enumerate(doc.paragraphs)
                     if p.text.strip().upper().startswith("INTRODUCTION")
                     or p.text.strip() == "Introduction")
    for i in range(abs_idx, intro_idx):
        p = doc.paragraphs[i]
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(3)
        for r in p.runs:
            _set_font(r, size=9)
    doc.paragraphs[abs_idx].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in doc.paragraphs[abs_idx].runs:
        r.font.bold = True

    # --- Insert continuous section break for two-column body ---
    body_sec = doc.add_section(WD_SECTION.CONTINUOUS)
    body_sec.page_width = Inches(8.5)
    body_sec.page_height = Inches(11)
    body_sec.top_margin = Inches(0.7)
    body_sec.bottom_margin = Inches(0.7)
    body_sec.left_margin = Inches(0.65)
    body_sec.right_margin = Inches(0.65)
    _set_columns(body_sec, 2, 360)

    # --- Heading formatting ---
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(3)
            _keep_with_next(p)
            for r in p.runs:
                _set_font(r, size=10, bold=True)
                r.text = r.text.upper()
        elif p.style.name == "Heading 2":
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(2)
            _keep_with_next(p)
            for r in p.runs:
                _set_font(r, size=9.5, bold=True, italic=True)

    # --- Center all tables and format cells ---
    for table in doc.tables:
        table.autofit = True
        _center_table(table)
        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                for cp in cell.paragraphs:
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cp.paragraph_format.space_after = Pt(0)
                    cp.paragraph_format.space_before = Pt(0)
                    for r in cp.runs:
                        _set_font(r, size=7.5, bold=(ri == 0))

    # --- Center all figures ---
    _center_inline_shapes(doc)

    # --- Center figure/table captions ---
    for p in doc.paragraphs:
        txt = p.text.strip()
        if (txt.startswith("Fig.") or txt.startswith("Figure")
                or txt.startswith("TABLE") or txt.startswith("Table")
                or txt.startswith("Algorithm")):
            if len(txt) > 15:  # likely a caption, not a reference
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    _set_font(r, size=8)

    # --- Insert REFERENCES heading ---
    first_ref = next((p for p in doc.paragraphs if p.text.strip().startswith("[1]")), None)
    if first_ref is not None:
        heading_p = first_ref.insert_paragraph_before("REFERENCES", style="Heading 1")
        heading_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in heading_p.runs:
            _set_font(r, size=10, bold=True)

    # --- Format reference entries ---
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt and txt[0] == "[" and "]" in txt[:5]:
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.8)
            p.paragraph_format.space_after = Pt(1)
            for r in p.runs:
                _set_font(r, size=7.5)

    # --- Document properties ---
    props = doc.core_properties
    props.title = ("CertiCut: Capacitated K-Way Partitioning with Anytime Bounds "
                   "for Sampling-Aware Circuit Cutting")
    props.subject = "IEEE conference paper"
    props.keywords = ("quantum circuit cutting; QPD; capacitated graph partitioning; "
                      "mixed-integer programming; solver-tolerance bounds")

    doc.save(DOCX)

    # --- Validation ---
    from zipfile import ZipFile
    with ZipFile(DOCX) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
    check = Document(DOCX)
    n_tables = len(check.tables)
    n_figs = len(check.inline_shapes)
    n_secs = len(check.sections)
    size = DOCX.stat().st_size
    print(f"Created {DOCX.name}: {size:,} bytes, "
          f"{n_tables} tables, {n_figs} figures, {n_secs} sections, "
          f"{len(media)} embedded media")
    if n_tables < 7:
        raise RuntimeError(f"Expected >=7 tables, got {n_tables}")
    if n_figs != 4:
        raise RuntimeError(f"Expected 4 figures, got {n_figs}")
    if size < 100_000:
        raise RuntimeError("File suspiciously small")


if __name__ == "__main__":
    format_document()
